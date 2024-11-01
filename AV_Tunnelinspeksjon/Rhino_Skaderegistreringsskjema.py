#! python3
# -*- coding: UTF-8 -*-
# r: python-docx, pandas, openpyxl

import System
import System.Windows.Forms as forms

import os

from docx import Document
from docx.shared import Cm
import locale
locale.setlocale(locale.LC_ALL, 'nb_NO')
import pandas as pd
import openpyxl

def createPage(document, regData, image_folder, tunnelinfo={'tunnel_name': 'Eksempeltunnel'}):
    """ INPUT: Referanse til Word-dokument, dict med data som skal inn i rapport, mappe hvor bildene er registrert. 
    Navn på tunnel i dict-format. (tenkt muligheter for utvidelse av funksjonalitet.. )
    OUTPUT: Ingen - Legger til ny side i referansedokument. 
    """
    try:
        document.add_heading('Skaderegistreringsskjema', 1)
        document.add_paragraph('Tunnelnavn: {}'.format(tunnelinfo['tunnel_name']))

        table1 = document.add_table(rows=3, cols=4)
        table1.style = document.styles['Light Grid']
        cell1 = table1.cell(0,3).merge(table1.cell(2,3))
        hdr_cells = table1.rows[0].cells
        hdr_cells[0].text = 'Retning:'
        hdr_cells[1].text = 'Pel'
        hdr_cells[2].text = 'Plassering:'
        body_cells = table1.rows[1].cells
        body_cells[0].text = regData['Retning']
        body_cells[1].text = str(regData['Pel'])

        paragraphX = cell1.paragraphs[0]
        run = paragraphX.add_run()

        if regData['Pos'] == 0:
            body_cells[2].text = 'Heng'
            run.add_picture(r'S:\Felles\SamferdselInfrastruktur\Rhino_GEO_Toolbar\plasseringer\heng.png', width=Cm(5))

        elif regData['Pos'] < 0:
            body_cells[2].text = 'Venstre side'
            run.add_picture(r'S:\Felles\SamferdselInfrastruktur\Rhino_GEO_Toolbar\plasseringer\vs_vegg.png', width=Cm(5))

        elif regData['Pos'] > 0:
            body_cells[2].text = 'Høyre side'
            run.add_picture(r'S:\Felles\SamferdselInfrastruktur\Rhino_GEO_Toolbar\plasseringer\hs_vegg.png', width=Cm(5))

        else:
            body_cells[2].text = 'Ukjent plassering'
            run.add_picture(r'S:\Felles\SamferdselInfrastruktur\Rhino_GEO_Toolbar\plasseringer\ukjent2.png', width=Cm(5))


        document.add_paragraph('', style='Normal')

        table2 = document.add_table(rows=2, cols=4)
        table2.style = document.styles['Light Grid']
        hdr_cells = table2.rows[0].cells
        hdr_cells[0].text = 'Registrering:'
        hdr_cells[1].text = 'Nedfall'
        hdr_cells[2].text = 'Plassforhold:'
        hdr_cells[3].text = 'Tiltak / Skadegrad'

        body_cells = table2.rows[1].cells
        body_cells[0].text = regData['Symbol']
        body_cells[1].text = regData['Nedfall']
        body_cells[2].text = regData['Plassforhold']
        body_cells[3].text = str(regData['Tiltak'])

        document.add_paragraph('', style='Normal')
        table3 = document.add_table(rows=3, cols=1)
        table3.style = document.styles['Table Grid']
        hdr_cell = table3.rows[0].cells
        hdr_cell[0].text = 'Beskrivelse:'
        body_cell = table3.rows[1].cells
        body_cell[0].text = regData['Kommentar']

        image_cell = table3.rows[2].cells
        image_cell[0]

        paragraphY = image_cell[0].paragraphs[0]
        run2 = paragraphY.add_run()
        run2.add_picture('{}\\{}'.format(image_folder, regData['Bildenavn']), width=Cm(14.89))

        document.add_page_break()

    except:
        print('Excel-arket er ikke satt opp rett')


def create_bilag3(tunnel_name, excel_fil, image_folder):
    """ INTPUT: Tunnelnavn-> Str, Excel fil med registreringer (sjekk referansefil for kolonneoverskrifter, mappe med bilder. 
    Fil blir lagret i samme mappe som bildene ligger med navn "Skaderegistreringer.docx"
    OUTPUT: Navn på fil dersom det lykkes. 
    """

    document = Document()
    
    df = pd.read_excel(excel_fil, header=0)
    df_dict = df.to_dict(orient="index")
    print(df_dict[0])
    
    for key in df_dict:
        createPage(document, df_dict[key], image_folder, {'tunnel_name': tunnel_name})
    
    docName = os.path.join(image_folder, '{}.docx'.format(tunnel_name))
    
    document.save(docName)
    print(docName, 'OK')
    return docName

def get_file():
    dialog = forms.OpenFileDialog()
    dialog.Filter = "Excel Files (*.xlsx)|*.xlsx"
    dialog.Title = "Velg Excel-fil"

    if dialog.ShowDialog() == forms.DialogResult.OK:
        return dialog.FileName
    return None

def get_directory():
    dialog = forms.FolderBrowserDialog()
    dialog.Description = "Velg mappe med jpg-bilder"
    dialog.ShowNewFolderButton = True

    if dialog.ShowDialog() == forms.DialogResult.OK:
        return dialog.SelectedPath
    return None

if __name__ == '__main__':
    excelFil = get_file()
    bilde_mappe = get_directory()
    wordFilePath = os.path.join(bilde_mappe,'Skaderegistreringer.docx')
    
    if os.path.isdir(bilde_mappe) and os.path.isfile(excelFil):
        create_bilag3('Skaderegistreringer',  excelFil, bilde_mappe)

    else: 
        pass